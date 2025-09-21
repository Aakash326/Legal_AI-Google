import re
import docx
from typing import List, Dict, Any, Tuple
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

class TextProcessor:
    def __init__(self):
        # Common legal document patterns
        self.section_patterns = [
            r'^\d+\.\s+',  # 1. Section
            r'^\(\d+\)\s+',  # (1) Section
            r'^[A-Z]\.\s+',  # A. Section
            r'^\([a-z]\)\s+',  # (a) Subsection
            r'^Article\s+\d+',  # Article 1
            r'^Section\s+\d+',  # Section 1
            r'^Chapter\s+\d+',  # Chapter 1
        ]
        
        # Legal clause indicators
        self.clause_indicators = [
            # Payment and Financial Terms
            'payment', 'fee', 'cost', 'charge', 'amount', 'price', 'invoice', 'billing', 'refund', 'penalty',
            'interest', 'late fee', 'deposit', 'installment', 'due date', 'payable', 'compensation',
            
            # Termination and Duration
            'termination', 'terminate', 'end', 'expir', 'cancel', 'dissolution', 'breach', 'default',
            'notice period', 'effective date', 'term', 'duration', 'renewal', 'extension',
            
            # Liability and Risk
            'liability', 'liable', 'responsible', 'damages', 'loss', 'injury', 'harm', 'negligence',
            'fault', 'limitation', 'exclusion', 'cap', 'consequential', 'indirect',
            
            # Indemnification
            'indemnif', 'hold harmless', 'defend', 'reimburse', 'compensate', 'make whole',
            
            # Confidentiality and Privacy
            'confidential', 'proprietary', 'non-disclosure', 'private', 'secret', 'privileged',
            'data protection', 'privacy', 'personal information', 'trade secret',
            
            # Intellectual Property
            'intellectual property', 'copyright', 'trademark', 'patent', 'trade mark', 'ip',
            'proprietary rights', 'license', 'ownership', 'derivative work',
            
            # Dispute Resolution
            'dispute', 'arbitration', 'litigation', 'court', 'mediation', 'resolution',
            'claim', 'action', 'proceeding', 'lawsuit', 'legal action',
            
            # Governing Law
            'governing law', 'jurisdiction', 'applicable law', 'venue', 'forum', 'choice of law',
            
            # Force Majeure
            'force majeure', 'act of god', 'unforeseeable', 'beyond control', 'natural disaster',
            
            # Amendments and Changes
            'amendment', 'modification', 'change', 'alter', 'revise', 'update', 'written consent',
            
            # Severability
            'severability', 'invalid', 'unenforceable', 'void', 'separate', 'remainder',
            
            # General Legal Terms
            'agreement', 'contract', 'party', 'parties', 'obligation', 'duty', 'right', 'warrant',
            'represent', 'covenant', 'undertake', 'bind', 'enforce', 'comply', 'violation',
            'breach', 'remedy', 'waiver', 'consent', 'approval', 'notice', 'delivery',
            
            # Performance and Delivery
            'performance', 'deliver', 'service', 'work', 'completion', 'milestone', 'deadline',
            'specification', 'standard', 'quality', 'acceptance', 'rejection',
            
            # Employment/Service Terms
            'employment', 'employee', 'contractor', 'service provider', 'worker', 'staff',
            'benefits', 'vacation', 'leave', 'resignation', 'dismissal', 'non-compete',
            
            # Data and Technology
            'data', 'software', 'system', 'technology', 'database', 'security', 'backup',
            'maintenance', 'support', 'update', 'upgrade', 'integration'
        ]
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            text = '\n'.join(text_parts)
            
            # Basic metadata
            metadata = {
                "extraction_method": "python-docx",
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "word_count": len(text.split()) if text else 0
            }
            
            # Try to get document properties
            try:
                if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
                if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
                if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                    metadata["created"] = str(doc.core_properties.created)
            except Exception as e:
                logger.warning(f"Could not extract document properties: {str(e)}")
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Could not extract text from DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "extraction_method": "plain_text",
                "word_count": len(text.split()) if text else 0,
                "character_count": len(text),
                "line_count": len(text.split('\n')) if text else 0
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                metadata = {
                    "extraction_method": "plain_text_latin1",
                    "word_count": len(text.split()) if text else 0,
                    "character_count": len(text),
                    "line_count": len(text.split('\n')) if text else 0
                }
                
                return text, metadata
                
            except Exception as e:
                logger.error(f"TXT extraction failed: {str(e)}")
                raise ValueError(f"Could not extract text from TXT: {str(e)}")
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            raise ValueError(f"Could not extract text from TXT: {str(e)}")
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common document artifacts
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\[Page \d+\]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        
        # Remove excessive blank lines
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    def chunk_text(self, text: str, max_chunk_size: int = 2000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks for AI processing"""
        if not text:
            return []
        
        # First try to split by sections
        sections = self._split_by_sections(text)
        
        chunks = []
        current_chunk = ""
        
        for section in sections:
            # If section is small enough, add it to current chunk
            if len(current_chunk) + len(section) <= max_chunk_size:
                current_chunk += "\n\n" + section if current_chunk else section
            else:
                # Save current chunk if it exists
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # If section is too large, split it further
                if len(section) > max_chunk_size:
                    section_chunks = self._split_large_section(section, max_chunk_size, overlap)
                    chunks.extend(section_chunks)
                    current_chunk = ""
                else:
                    current_chunk = section
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return [chunk for chunk in chunks if chunk.strip()]
    
    def _split_by_sections(self, text: str) -> List[str]:
        """Split text by legal document sections"""
        lines = text.split('\n')
        sections = []
        current_section = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_section:
                    current_section.append("")
                continue
            
            # Check if this line starts a new section
            is_section_start = any(re.match(pattern, line) for pattern in self.section_patterns)
            
            if is_section_start and current_section:
                # Save current section and start new one
                sections.append('\n'.join(current_section))
                current_section = [line]
            else:
                current_section.append(line)
        
        # Add the last section
        if current_section:
            sections.append('\n'.join(current_section))
        
        return [section for section in sections if section.strip()]
    
    def _split_large_section(self, section: str, max_size: int, overlap: int) -> List[str]:
        """Split a large section into smaller chunks with overlap"""
        words = section.split()
        chunks = []
        
        i = 0
        while i < len(words):
            chunk_words = []
            chunk_size = 0
            
            # Build chunk up to max_size
            while i < len(words) and chunk_size < max_size:
                word = words[i]
                if chunk_size + len(word) + 1 <= max_size:  # +1 for space
                    chunk_words.append(word)
                    chunk_size += len(word) + 1
                    i += 1
                else:
                    break
            
            if chunk_words:
                chunks.append(' '.join(chunk_words))
                
                # Calculate overlap for next chunk
                if i < len(words):  # More words to process
                    overlap_words = min(overlap // 10, len(chunk_words) // 4)  # Rough estimate
                    i -= overlap_words
        
        return chunks
    
    def identify_potential_clauses(self, text: str) -> List[Dict[str, Any]]:
        """Identify potential legal clauses in text"""
        clauses = []
        
        # Split text into sentences/paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        # Also split by sentences for better clause detection
        sentences = []
        for paragraph in paragraphs:
            # Split by periods, but keep sentences that are likely legal clauses
            paragraph_sentences = re.split(r'\.(?=\s+[A-Z])', paragraph)
            for sentence in paragraph_sentences:
                sentence = sentence.strip()
                if len(sentence) > 50:  # Only consider substantial sentences
                    sentences.append(sentence)
        
        # Analyze paragraphs first
        for i, paragraph in enumerate(paragraphs):
            if len(paragraph) > 100:  # Only analyze substantial paragraphs
                clause_info = self._analyze_paragraph_for_clauses(paragraph, i)
                if clause_info:
                    clauses.extend(clause_info)
        
        # Analyze sentences for additional clauses
        for i, sentence in enumerate(sentences[:50]):  # Limit to first 50 sentences
            if len(sentence) > 100:  # Only analyze substantial sentences
                clause_info = self._analyze_sentence_for_clauses(sentence, i + len(paragraphs))
                if clause_info:
                    clauses.extend(clause_info)
        
        # Remove duplicates and sort by importance
        unique_clauses = []
        seen_texts = set()
        
        for clause in clauses:
            clause_text_lower = clause['text'].lower()
            # Create a signature for similarity checking
            signature = ' '.join(sorted(clause_text_lower.split()[:10]))  # First 10 words sorted
            
            if signature not in seen_texts:
                seen_texts.add(signature)
                unique_clauses.append(clause)
        
        # Sort by importance score and return top 20
        unique_clauses.sort(key=lambda x: x.get('importance_score', 0), reverse=True)
        return unique_clauses[:20]
    
    def _analyze_paragraph_for_clauses(self, paragraph: str, paragraph_index: int) -> List[Dict[str, Any]]:
        """Analyze a paragraph for legal clause indicators"""
        clauses = []
        
        # Look for clause indicators
        paragraph_lower = paragraph.lower()
        
        # Find ALL matching indicators, not just the first one
        matching_indicators = []
        for indicator in self.clause_indicators:
            if indicator in paragraph_lower:
                matching_indicators.append(indicator)
        
        if matching_indicators:
            # Get the most prominent clause type
            clause_type = self._categorize_clause(paragraph_lower, matching_indicators[0])
            
            clauses.append({
                "text": paragraph,
                "paragraph_index": paragraph_index,
                "indicators": matching_indicators,  # Store all matching indicators
                "potential_type": clause_type,
                "importance_score": self._calculate_importance_score(paragraph),
                "source": "paragraph"
            })
        
        return clauses
    
    def _analyze_sentence_for_clauses(self, sentence: str, sentence_index: int) -> List[Dict[str, Any]]:
        """Analyze a sentence for legal clause indicators"""
        clauses = []
        
        # Look for clause indicators
        sentence_lower = sentence.lower()
        
        # Find matching indicators
        matching_indicators = []
        for indicator in self.clause_indicators:
            if indicator in sentence_lower:
                matching_indicators.append(indicator)
        
        if matching_indicators:
            # Get the most prominent clause type
            clause_type = self._categorize_clause(sentence_lower, matching_indicators[0])
            
            clauses.append({
                "text": sentence,
                "sentence_index": sentence_index,
                "indicators": matching_indicators,
                "potential_type": clause_type,
                "importance_score": self._calculate_importance_score(sentence),
                "source": "sentence"
            })
        
        return clauses
    
    def _categorize_clause(self, paragraph_lower: str, indicator: str) -> str:
        """Categorize clause based on indicators"""
        if any(word in paragraph_lower for word in ['payment', 'fee', 'cost', 'charge', 'amount']):
            return "payment_terms"
        elif any(word in paragraph_lower for word in ['termination', 'terminate', 'end', 'expir']):
            return "termination"
        elif any(word in paragraph_lower for word in ['liability', 'liable', 'responsible', 'damages']):
            return "liability"
        elif any(word in paragraph_lower for word in ['confidential', 'proprietary', 'non-disclosure']):
            return "confidentiality"
        elif any(word in paragraph_lower for word in ['intellectual property', 'copyright', 'trademark']):
            return "intellectual_property"
        elif any(word in paragraph_lower for word in ['dispute', 'arbitration', 'litigation', 'court']):
            return "dispute_resolution"
        elif any(word in paragraph_lower for word in ['governing law', 'jurisdiction']):
            return "governing_law"
        elif any(word in paragraph_lower for word in ['amendment', 'modification', 'change']):
            return "amendment"
        else:
            return "other"
    
    def _calculate_importance_score(self, paragraph: str) -> float:
        """Calculate importance score for a paragraph"""
        score = 0.0
        paragraph_lower = paragraph.lower()
        
        # Length score (longer paragraphs might be more important)
        score += min(len(paragraph) / 1000, 1.0) * 0.3
        
        # Legal keyword density
        legal_keywords = ['shall', 'must', 'required', 'obligation', 'right', 'liability', 'agreement']
        keyword_count = sum(1 for keyword in legal_keywords if keyword in paragraph_lower)
        score += (keyword_count / len(legal_keywords)) * 0.4
        
        # Structural indicators (numbered sections, etc.)
        if any(re.match(pattern, paragraph) for pattern in self.section_patterns):
            score += 0.3
        
        return min(score, 1.0)
    
    def get_word_count(self, text: str) -> int:
        """Get word count of text"""
        return len(text.split()) if text else 0
    
    def get_character_count(self, text: str) -> int:
        """Get character count of text"""
        return len(text) if text else 0

# Global instance
text_processor = TextProcessor()